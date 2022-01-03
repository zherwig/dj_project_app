from django.http import HttpResponse
from django.shortcuts import render, get_object_or_404, redirect
from django.contrib.auth.models import User
from django.contrib.auth.decorators import login_required
from django.http import StreamingHttpResponse

from docx import Document
import datetime
import io

from projects.models import Project
from tasks.models import Task
from dashboards import applogic as dashboard_applogic

# Create your views here.
@login_required
def dashboard_list_view(request, *args, **kwargs):
    context = {
        "topics": dashboard_applogic.get_topics_with_projects_and_open_tasks(),
        "user":request.user,
    }
    
    return render(request, 'dashboard_list.html', context) 


@login_required
def dashboard_actions_view(request, *args, **kwargs): 
    context = {
        "quick_action_task": Task.objects.get(title="Random stuff"),
        "task_batches": [
            {
                'name':'Overdue',
                'tasks': dashboard_applogic.get_staff_actions_per_date_range(-900, -1, "zjef")
            },
            {
                'name':'Today',
                'tasks': dashboard_applogic.get_staff_actions_per_date_range(0, 0, "zjef")
            },
            {
                'name':'Tomorrow',
                'tasks': dashboard_applogic.get_staff_actions_per_date_range(1, 1, "zjef")
            },
            {
                'name':'Rest of the week',
                'tasks': dashboard_applogic.get_staff_actions_per_date_range(2, 7, "zjef")
            },
            {
                'name':'Rest of the month',
                'tasks': dashboard_applogic.get_staff_actions_per_date_range(8, 30, "zjef")
            },
        ],
        "user":request.user,
    }
    return render(request, 'dashboard_todos.html', context) 

@login_required
def dashboard_actions_fix_overdues(request, *args, **kwargs):   
    overdue_actions = dashboard_applogic.get_actions_per_date_range(-900, -1)
    for overdue_action in overdue_actions:
        dashboard_applogic.move_action_to_today(overdue_action.id)
    return redirect("/today")

def dashboard_today_actions_push_all(request, *args, **kwargs):
    today_actions = dashboard_applogic.get_staff_actions_per_date_range(0, 0, "zjef")
    for today_action in today_actions:
        action_result = dashboard_applogic.move_action_to_tomorrow(today_action.id)
        if not action_result:
            return HttpResponse(f'Exception: {action_result}')
    return redirect("/today")

def dashboard_today_actions_unmute_all(request, *args, **kwargs):
    today_actions = dashboard_applogic.get_staff_actions_per_date_range(0, 0, "zjef")
    for today_action in today_actions:
        action_result = dashboard_applogic.unmute_action(today_action.id)
        if not action_result:
            return HttpResponse(f'Exception: {action_result}')
    return redirect("/today")

@login_required
def dashboard_today_view(request, *args, **kwargs):
    context = {
        "quick_action_task": Task.objects.get(title="Random stuff"),
        "todays_actions": dashboard_applogic.get_staff_actions_per_date_range(0, 0, "zjef"),
        "user":request.user,
    }
    return render(request, 'dashboard_today.html', context) 

@login_required
def docx_report(request, *args, **kwargs):
    all_topics = dashboard_applogic.get_topics_with_projects_and_open_tasks()
    all_topics.reverse()
    completed_topics = dashboard_applogic.get_completed_tasks()
    context = {
        "completed": completed_topics,
        "topics": all_topics,
        "user":request.user
    }
    
    return render(request, 'report_prep.html', context)

@login_required
def download_docx(request, *args, **kwargs):
    completed_topics = dashboard_applogic.get_completed_tasks()
    all_topics = dashboard_applogic.get_topics_with_projects_and_open_tasks()
    all_topics.reverse()
    # create an empty document object
    document = Document()
    
    #contents
    document.add_heading('Weekly project progress report', 0)
    document.add_heading('Completed items this week', level=1)
    for completed_topic in completed_topics:
        completed_topic_p = document.add_paragraph(
                            '', style='List Bullet'
                        )
        description_header = completed_topic_p.add_run(f'{completed_topic.title}')
        description_header.bold = True
        task_run = completed_topic_p.add_run()
        task_run.add_break()
        description_header = completed_topic_p.add_run("Description: ")
        description_header.italic = True
        completed_topic_p.add_run(f"{completed_topic.detail}")
        task_run = completed_topic_p.add_run()
        task_run.add_break()
        progress_header = completed_topic_p.add_run("Project: ")
        progress_header.italic = True
        completed_topic_p.add_run(f"{completed_topic.project.title}")
    document.add_heading('Open items', level=1)
    for topic in all_topics:
        document.add_heading(f"{topic['topic']}", level=3)
        for project in topic['projects']:
            if project.excludeFromReports == False:
                document.add_paragraph(
                f'{project.title}', style='List Number'
                )
                for task in project.tasks:
                    if task.excludeFromReports == False:
                        task_p = document.add_paragraph(
                            '', style='List Bullet 2'
                        )
                        task_p.add_run(f'{task.title}').bold = True
                        task_run = task_p.add_run()
                        task_run.add_break()
                        description_header = task_p.add_run("Description: ")
                        description_header.italic = True
                        task_p.add_run(f"{task.detail}")
                        task_run = task_p.add_run()
                        task_run.add_break()
                        progress_header = task_p.add_run("Progress: ")
                        progress_header.italic = True
                        task_p.add_run(f"{task.taskProgress}")
                        task_run = task_p.add_run()
                        task_run.add_break()
                        duedate_header = task_p.add_run("Due Date:")
                        duedate_header.italic = True
                        task_p.add_run(f"{task.duedate}")
                
    # save document info
    buffer = io.BytesIO()
    document.save(buffer)  # save your memory stream
    buffer.seek(0)  # rewind the stream

    # put them to streaming content response 
    # within docx content_type
    response = StreamingHttpResponse(
        streaming_content=buffer,  # use the stream's content
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingm'
    )

    response['Content-Disposition'] = 'attachment;filename=Weekly_Report.docx'
    response["Content-Encoding"] = 'UTF-8'

    return response
